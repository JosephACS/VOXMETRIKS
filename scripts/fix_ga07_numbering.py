"""
Numeracion continua ENTREGABLE GA07:

  Parte I:  1 (SDD) -> 2 (Constitucion) -> 3-13 (modulos 001-011) -> 14 (matriz)
  Parte II: 15-19 (diagramas UML, sin reiniciar en 1 ni repetir 12)

Uso: python fix_ga07_numbering.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

GA07 = Path(__file__).resolve().parent
BLACK = RGBColor(0, 0, 0)
NUM_PREFIX = re.compile(r"^(?:\d+\.)+\d*\s+")

MATRIZ_NUM = 14
P2_START = 15

METODO_TITLES = [
    (6, "Tabla modelo — partes del sistema"),
    (5, "Estructura de carpetas del proyecto"),
    (4, "Ejemplo login — tabla Elemento / Ejemplo"),
    (3, "Plantilla de trabajo"),
    (2, "Estructura de una buena especificacion"),
    (1, "Fases SDD"),
]

P2_SECTIONS = [
    ("casos de uso", "diagrama", "Diagrama de casos de uso"),
    ("componentes", None, "Diagrama de componentes"),
    ("arquitectura", None, "Diagrama de arquitectura"),
    ("flujo", "diagrama", "Diagrama flujo ELT"),
    ("trazabilidad uml", None, "Trazabilidad UML y specs"),
]


def strip_num(text: str) -> str:
    t = text.strip()
    while True:
        n = NUM_PREFIX.sub("", t, count=1)
        if n == t:
            return t.strip()
        t = n


def set_heading_text(para, new_text: str) -> None:
    for run in list(para.runs):
        run.text = ""
    run = para.runs[0] if para.runs else para.add_run()
    run.text = new_text
    run.font.color.rgb = BLACK
    run.font.name = "Calibri"
    run.font.bold = True


def detect_metodo_sub(raw: str) -> tuple[int, str] | None:
    m = re.match(r"^1\.(\d+)\s", raw)
    if m:
        i = int(m.group(1))
        for num, title in METODO_TITLES:
            if num == i:
                return num, title
    low = raw.lower()
    for num, title in METODO_TITLES:
        if title.lower() in low:
            return num, title
    return None


def classify_p2(raw: str) -> str | None:
    low = raw.lower()
    for a, b, title in P2_SECTIONS:
        if b:
            if a in low and b in low:
                return title
        elif a in low:
            return title
    if "elt" in low and "diagrama" in low:
        return "Diagrama flujo ELT"
    return None


def renumber_document(doc: Document) -> None:
    zone = "start"
    in_parte2 = False
    const_n = 0
    const_sub = 0
    mod_major = 0
    mod_minor = 0
    mod_sub = 0
    p2_idx = 0

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        if not style.startswith("Heading"):
            continue
        level = int(style.split()[-1])
        raw = para.text.strip()
        if not raw:
            continue

        # --- Heading 1: zonas principales ---
        if level == 1:
            in_parte2 = "Parte II" in raw
            if raw.startswith("Anexo"):
                zone = "anexo"
            elif "Parte I" in raw and "Parte II" not in raw:
                zone = "parte1"
                in_parte2 = False
            elif in_parte2:
                zone = "parte2"
                p2_idx = 0
            elif "Resumen ejecutivo" in raw:
                in_parte2 = False
            continue

        # --- Parte II: 15, 16, 17, 18, 19 ---
        if in_parte2 and level == 2:
            title = classify_p2(raw) or strip_num(raw)
            for _, _, t in P2_SECTIONS:
                if t.lower() in title.lower() or title.lower() in t.lower():
                    title = t
                    break
            p2_idx += 1
            num = P2_START + p2_idx - 1
            set_heading_text(para, f"{num}. {title}")
            continue

        # --- Metodologia 1.x ---
        if zone in ("parte1", "start") and level == 2 and "Metodologia" in raw:
            zone = "metodo"
            set_heading_text(para, "1. Metodologia Spec Driven Development")
            continue

        if zone == "metodo" and level == 3:
            hit = detect_metodo_sub(raw)
            if hit:
                i, title = hit
                set_heading_text(para, f"1.{i} {title}")
            continue

        # --- Modulo N (3-13) ---
        if level == 2 and ("Especificacion operativa" in raw or re.search(r"Modulo 0\d{2}-", raw)):
            zone = "spec"
            in_parte2 = False
            mod_minor = 0
            mod_sub = 0
            m = re.search(r"Modulo (0\d{2})-", raw)
            mod_major = int(m.group(1)) + 2 if m else mod_major + 1
            folder_m = re.search(r"(0\d{2}-[\w-]+)", raw)
            folder = folder_m.group(1) if folder_m else "???"
            set_heading_text(para, f"{mod_major}. Especificacion operativa — Modulo {folder}")
            continue

        # --- Constitucion 2.x ---
        if level == 2 and "Constitucion del Proyecto" in raw:
            zone = "const"
            in_parte2 = False
            const_n = 0
            const_sub = 0
            set_heading_text(para, "2. Archivo de Constitucion del Proyecto")
            continue

        if zone == "const" and level == 2:
            const_n += 1
            const_sub = 0
            set_heading_text(para, f"2.{const_n} {strip_num(raw)}")
            continue

        if zone == "const" and level in (3, 4):
            const_sub += 1
            set_heading_text(para, f"2.{const_n}.{const_sub} {strip_num(raw)}")
            continue

        # --- Matriz 14 (antes de spec subsecciones sueltas) ---
        if level == 2 and "Matriz maestra" in raw:
            zone = "matrix"
            in_parte2 = False
            set_heading_text(para, f"{MATRIZ_NUM}. Matriz maestra de trazabilidad — Resumen por modulo")
            continue

        # --- Subsecciones modulo N.x ---
        if zone == "spec" and level == 2:
            mod_minor += 1
            mod_sub = 0
            set_heading_text(para, f"{mod_major}.{mod_minor} {strip_num(raw)}")
            continue

        if zone == "spec" and level in (3, 4):
            mod_sub += 1
            set_heading_text(para, f"{mod_major}.{mod_minor}.{mod_sub} {strip_num(raw)}")
            continue


def verify_document(doc: Document) -> list[str]:
    errors: list[str] = []
    h2_nums: list[tuple[str, int]] = []

    for para in doc.paragraphs:
        if not (para.style and para.style.name == "Heading 2"):
            continue
        raw = para.text.strip()
        m = re.match(r"^(\d+)\.\s", raw)
        if m:
            h2_nums.append((raw[:70], int(m.group(1))))

    # Ultimo de Parte I antes de Parte II: debe ser 14
    parte2_i = next((i for i, p in enumerate(doc.paragraphs) if p.style and p.style.name == "Heading 1" and "Parte II" in p.text), None)
    if parte2_i is None:
        errors.append("No se encontro Parte II")
        return errors

    pre_p2 = [
        p.text.strip() for p in doc.paragraphs[:parte2_i]
        if p.style and p.style.name == "Heading 2" and re.match(r"^\d+\.", p.text.strip())
    ]
    post_p2 = [
        p.text.strip() for p in doc.paragraphs[parte2_i:]
        if p.style and p.style.name == "Heading 2" and re.match(r"^\d+\.", p.text.strip())
    ]

    if pre_p2:
        last = pre_p2[-1]
        if not last.startswith(f"{MATRIZ_NUM}."):
            errors.append(f"Parte I debe terminar en {MATRIZ_NUM}. Matriz; ultimo H2: {last[:60]}")

    if post_p2:
        first = post_p2[0]
        if not first.startswith(f"{P2_START}."):
            errors.append(f"Parte II debe empezar en {P2_START}.; primer H2: {first[:60]}")
        for i, t in enumerate(post_p2[:5]):
            expected = P2_START + i
            m = re.match(r"^(\d+)\.", t)
            if not m or int(m.group(1)) != expected:
                errors.append(f"Parte II seccion {i+1} debe ser {expected}.; es: {t[:60]}")

    # No debe haber 12.x en Parte II
    for t in post_p2:
        if re.match(r"^12\.", t):
            errors.append(f"Parte II aun tiene numeracion 12.x: {t[:60]}")

    return errors


def main() -> None:
    src = None
    for name in (
        "ENTREGABLE-GA07-Voxmetriks-FINAL.docx",
        "ENTREGABLE-GA07-Voxmetriks.docx",
        "ENTREGABLE-GA07-Voxmetriks-NUMERADO.docx",
    ):
        p = GA07 / name
        if p.exists():
            src = p
            break
    if not src:
        sys.exit("No hay docx en GA07")

    doc = Document(src)
    renumber_document(doc)
    errors = verify_document(doc)
    if errors:
        print("ERRORES DE VERIFICACION:", file=sys.stderr)
        for e in errors:
            print(f"  - {e}", file=sys.stderr)
        sys.exit(1)

    out = GA07 / "ENTREGABLE-GA07-Voxmetriks.docx"
    try:
        doc.save(out)
        print(f"OK: {out}")
    except PermissionError:
        out = GA07 / "ENTREGABLE-GA07-Voxmetriks-NUMERADO.docx"
        doc.save(out)
        print(f"Word abierto — guardado: {out}")

    # Resumen transicion Parte I -> II
    parte2_i = next(i for i, p in enumerate(doc.paragraphs) if p.style and "Heading 1" in p.style.name and "Parte II" in p.text)
    pre = [p.text for p in doc.paragraphs[:parte2_i] if p.style and p.style.name == "Heading 2" and re.match(r"^1[34]\.", p.text.strip())]
    post = [p.text for p in doc.paragraphs[parte2_i:parte2_i + 30] if p.style and p.style.name == "Heading 2" and re.match(r"^1[5-9]\.", p.text.strip())]
    print("Cierre Parte I:", pre[-1][:70] if pre else "?")
    print("Inicio Parte II:", post[0][:70] if post else "?")


if __name__ == "__main__":
    main()
